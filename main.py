import os
import json
import logging
from datetime import datetime
from typing import Dict, List

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ContextTypes,
    ConversationHandler,
    filters,
)

TELEGRAM_TOKEN = ""  
AUDIO_DIR = "audio"
TEMPLATE_PATH = "templates/template.docx"
DATA_FILE = "data.json"


os.makedirs(AUDIO_DIR, exist_ok=True)
os.makedirs("templates", exist_ok=True)


logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

SELECTING_RECORD, EDITING_SUMMARY = range(2)


whisper_model = None
summarizer_pipeline = None

if not os.path.exists(TEMPLATE_PATH):
    from docx import Document
    doc = Document()
    doc.add_heading('Отчет по аудиозаписи', 0)
    doc.add_paragraph(f'Дата создания: [DATE]')
    doc.add_paragraph('Содержание:')
    doc.add_paragraph('[SUMMARY_TEXT]')
    doc.save(TEMPLATE_PATH)


def load_data() -> Dict:
    if not os.path.exists(DATA_FILE):
        return {"users": {}}
    try:
        with open(DATA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        return {"users": {}}

def save_data(data: Dict):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def get_or_create_user(tg_user_id: int, username: str) -> Dict:
    data = load_data()
    if str(tg_user_id) not in data["users"]:
        data["users"][str(tg_user_id)] = {
            "username": username,
            "records": []
        }
        save_data(data)
    return data["users"][str(tg_user_id)]

def save_voice_record(user_id: int, audio_path: str, raw_text: str, summary_text: str) -> int:
    data = load_data()
    record = {
        "id": int(datetime.now().timestamp()),
        "audio_path": audio_path,
        "raw_text": raw_text,
        "summary_text": summary_text,
        "created_at": datetime.now().isoformat()
    }
    data["users"][str(user_id)]["records"].append(record)
    save_data(data)
    return record["id"]

def get_user_records(user_id: int) -> List[Dict]:
    data = load_data()
    return data["users"].get(str(user_id), {}).get("records", [])

def get_record_by_id(user_id: int, record_id: int) -> Dict:
    records = get_user_records(user_id)
    for rec in records:
        if rec["id"] == record_id:
            return rec
    return None

def update_summary_text(user_id: int, record_id: int, new_summary: str):
    data = load_data()
    for rec in data["users"][str(user_id)]["records"]:
        if rec["id"] == record_id:
            rec["summary_text"] = new_summary
            break
    save_data(data)


def load_models():
    global whisper_model, summarizer_pipeline
    try:
        if whisper_model is None:
            import whisper
            logger.info("Загрузка модели Whisper...")
            whisper_model = whisper.load_model("tiny", device="cpu")
        if summarizer_pipeline is None:
            from transformers import pipeline
            logger.info("Загрузка модели суммаризации...")
            summarizer_pipeline = pipeline(
                "summarization",
                model="facebook/bart-large-cnn",
                device=-1,
                truncation=True
            )
    except Exception as e:
        logger.error(f"Ошибка загрузки моделей: {e}")
    return whisper_model, summarizer_pipeline

def transcribe_audio(audio_path: str) -> str:
    try:
        whisper_model, _ = load_models()
        logger.info(f"Распознавание аудио: {audio_path}")
        result = whisper_model.transcribe(audio_path, language="ru")
        return result["text"]
    except Exception as e:
        logger.error(f"Ошибка распознавания: {e}")
        return "Не удалось понять тебя, родной"

def summarize_text(text: str) -> str:
    try:
        _, summarizer_pipeline = load_models()
        if len(text) < 50:
            return text
            
        logger.info("Суммаризация текста...")
        summary = summarizer_pipeline(
            text,
            max_length=150,
            min_length=30,
            do_sample=False
        )[0]['summary_text']
        return summary
    except Exception as e:
        logger.error(f"Ошибка суммаризации: {e}")
        return text[:100] + "..." if len(text) > 100 else text

from docx import Document

def generate_docx(summary_text: str, output_path: str):
    try:
        doc = Document(TEMPLATE_PATH)
        for paragraph in doc.paragraphs:
            if "[SUMMARY_TEXT]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[SUMMARY_TEXT]", summary_text)
            if "[DATE]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[DATE]", datetime.now().strftime("%d.%m.%Y %H:%M"))
        doc.save(output_path)
    except Exception as e:
        logger.error(f"Ошибка создания DOCX: {e}")
        # Создаем простой документ как запасной вариант
        doc = Document()
        doc.add_heading('Отчет', 0)
        doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph(summary_text)
        doc.save(output_path)


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    get_or_create_user(user.id, user.username)
    
    keyboard = [
        [InlineKeyboardButton(" Записать голос", callback_data="record_voice")],
        [InlineKeyboardButton(" Создать документ", callback_data="create_document")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    if update.message:
        await update.message.reply_text(
            f"Привет, {user.first_name}! \n\n"
            "Я бот для записи голоса, распознавания и создания отчётов.\n\n"
            "Выбери действие:",
            reply_markup=reply_markup
        )
    else:
        await update.callback_query.edit_message_text(
            f"Привет, {user.first_name}! \n\n"
            "Я бот для записи голоса, распознавания и создания отчётов.\n\n"
            "Выбери действие:",
            reply_markup=reply_markup
        )

async def record_voice_button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.edit_message_text(" Отправьте аудиосообщение — я его услышу.")

async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    get_or_create_user(user.id, user.username)

    voice = update.message.voice
    file_id = voice.file_id
    file_ext = ".ogg"
    audio_path = os.path.join(AUDIO_DIR, f"{file_id}{file_ext}")

    try:
        file = await context.bot.get_file(file_id)
        await file.download_to_drive(audio_path)
        logger.info(f"Аудио сохранено: {audio_path}")

        await update.message.reply_text("Слушаю аудио...")

        raw_text = transcribe_audio(audio_path)
        summary_text = summarize_text(raw_text)

        save_voice_record(user.id, audio_path, raw_text, summary_text)

        await update.message.reply_text(
            f"Услышал тебя, родной!\n\n"
            f" **Содержание:**\n{summary_text}\n\n"
            f"Теперь ты можешь создать документ — нажми «Создать документ»."
        )
    except Exception as e:
        logger.error(f"Ошибка обработки голоса: {e}")
        await update.message.reply_text(" Произошла ошибка при обработке аудио.")

async def create_document_button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    user_id = query.from_user.id

    records = get_user_records(user_id)
    if not records:
        await query.edit_message_text("❌ У вас ещё нет записей.")
        return

    keyboard = []
    for rec in sorted(records, key=lambda r: r["created_at"], reverse=True)[:10]:  # ограничиваем 10 записями
        date_str = datetime.fromisoformat(rec['created_at']).strftime('%d.%m %H:%M')
        btn_text = f"{date_str} — {rec['summary_text'][:30]}..."
        keyboard.append([InlineKeyboardButton(btn_text, callback_data=f"select_{rec['id']}")])

    keyboard.append([InlineKeyboardButton("↩️ Назад", callback_data="start")])
    reply_markup = InlineKeyboardMarkup(keyboard)

    await query.edit_message_text(
        "📋 Выберите запись для создания документа:",
        reply_markup=reply_markup
    )

async def select_record(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    record_id = int(query.data.split("_")[1])
    user_id = query.from_user.id

    record = get_record_by_id(user_id, record_id)
    if not record:
        await query.edit_message_text("Ошибка: запись не найдена.")
        return ConversationHandler.END

    context.user_data["selected_record_id"] = record_id
    context.user_data["current_summary"] = record["summary_text"]

    await query.edit_message_text(
        f"Вы выбрали запись:\n\n"
        f" Исходный текст:\n{record['raw_text'][:100]}...\n\n"
        f" Сжатый текст (редактируемый):\n{record['summary_text']}\n\n"
        f" Теперь отправьте мне **новый текст**, который нужно вставить в документ.\n"
        f"(или отправьте /cancel, чтобы отменить)"
    )
    return EDITING_SUMMARY

async def edit_summary(update: Update, context: ContextTypes.DEFAULT_TYPE):
    new_summary = update.message.text.strip()
    record_id = context.user_data.get("selected_record_id")
    user_id = update.effective_user.id

    if not record_id:
        await update.message.reply_text("Произошла ошибка. Попробуйте снова.")
        return ConversationHandler.END

    update_summary_text(user_id, record_id, new_summary)
    context.user_data["current_summary"] = new_summary

    output_path = os.path.join(AUDIO_DIR, f"report_{record_id}.docx")
    generate_docx(new_summary, output_path)

    with open(output_path, "rb") as doc:
        await update.message.reply_document(doc, caption=" Ваш готовый документ!")

    context.user_data.clear()
    
    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("🏠 Главное меню", callback_data="start")]
    ])
    await update.message.reply_text(
        " Документ создан и отправлен!\n\n"
        "Чтобы сделать ещё один — нажмите «Создать документ».",
        reply_markup=keyboard
    )
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    await update.message.reply_text("✏️ Редактирование отменено.")
    return ConversationHandler.END

async def back_to_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await start(update, context)


def main():
    application = Application.builder().token(TELEGRAM_TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[CallbackQueryHandler(select_record, pattern=r"^select_\d+$")],
        states={
            EDITING_SUMMARY: [MessageHandler(filters.TEXT & ~filters.COMMAND, edit_summary)],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )

    application.add_handler(CommandHandler("start", start))
    application.add_handler(CallbackQueryHandler(record_voice_button, pattern="^record_voice$"))
    application.add_handler(CallbackQueryHandler(create_document_button, pattern="^create_document$"))
    application.add_handler(CallbackQueryHandler(back_to_start, pattern="^start$"))
    application.add_handler(MessageHandler(filters.VOICE, handle_voice))
    application.add_handler(conv_handler)

    logger.info("🚀 Бот запущен...")

    application.run_polling()

if __name__ == "__main__":
    main()